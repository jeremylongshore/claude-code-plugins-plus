#!/usr/bin/env python3
"""
build_standards_docx.py

Convert 5 canonical Standard of Truth Markdown files to DOCX teaching editions.
Handles headers, code blocks, tables, lists, and embeds diagrams.

Usage:
    python build_standards_docx.py

    Or use the smoke test script for validation:
    ./scripts/smoke_docx_build.sh

Features:
    - Converts Markdown to professionally formatted DOCX
    - Preserves heading hierarchy (H1, H2, H3)
    - Formats code blocks with monospace font (Courier New)
    - Converts Markdown tables to Word tables with borders
    - Embeds diagram PNG images (6 inches wide)
    - Includes Mermaid source code after each diagram
    - Handles bulleted and numbered lists
    - Professional styling (Calibri font, blue headings)

Outputs 5 DOCX files to 000-docs/:
    - 6767-c-DR-STND-claude-code-extensions-standard.docx
    - 6767-d-AT-STND-claude-code-extensions-schema.docx
    - 6767-e-WA-STND-extensions-validation-and-ci-gates.docx
    - 6767-f-AT-ARCH-plugin-scaffold-diagrams.docx (includes 3 diagrams)
    - 6767-g-AT-ARCH-skill-scaffold-diagrams.docx (includes 4 diagrams)

Dependencies:
    - python-docx==1.2.0 (install: pip install -r scripts/requirements.txt)

Diagram Embedding:
    - Automatically detects references: **See diagram**: [filename.mmd]
    - Embeds corresponding .png file
    - Includes Mermaid source as code block below image
"""

import re
import sys
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Project root
PROJECT_ROOT = Path(__file__).parent.parent
DOCS_DIR = PROJECT_ROOT / "000-docs"

# Standards to convert (5 canonical files)
STANDARDS = [
    "6767-c-DR-STND-claude-code-extensions-standard",
    "6767-d-AT-STND-claude-code-extensions-schema",
    "6767-e-WA-STND-extensions-validation-and-ci-gates",
    "6767-f-AT-ARCH-plugin-scaffold-diagrams",
    "6767-g-AT-ARCH-skill-scaffold-diagrams",
]


def setup_document_styles(doc: Document) -> None:
    """Configure document-wide styles for professional appearance."""
    # Normal style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)

    # Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0, 70, 127)

        if i == 1:
            heading_font.size = Pt(18)
        elif i == 2:
            heading_font.size = Pt(14)
        else:
            heading_font.size = Pt(12)


def add_table_borders(table) -> None:
    """Add borders to table for professional appearance."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def parse_markdown_table(lines: List[str]) -> Tuple[List[str], List[List[str]]]:
    """Parse a Markdown table into headers and rows."""
    # First line is headers
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]

    # Skip separator line (line 1)
    # Remaining lines are data rows
    rows = []
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)

    return headers, rows


def add_markdown_table(doc: Document, lines: List[str]) -> None:
    """Add a Markdown table to the document."""
    headers, rows = parse_markdown_table(lines)

    if not headers or not rows:
        return

    # Create table
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    add_table_borders(table)

    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        # Bold header text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(table.columns):
                table.rows[row_idx].cells[col_idx].text = cell_data


def find_diagram_files(standard_name: str) -> Dict[str, Tuple[str, str]]:
    """Find all diagram .png and .mmd files for a given standard.

    Returns dict: {reference_text: (png_path, mmd_path)}
    Example: {"6767-f-diagram-1-plugin-anatomy": ("/path/to.png", "/path/to.mmd")}
    """
    diagrams = {}

    # Find all .png files for this standard
    pattern = f"{standard_name.split('-')[0]}-{standard_name.split('-')[1]}-diagram-*.png"
    for png_file in DOCS_DIR.glob(pattern):
        # Extract diagram reference (e.g., "6767-f-diagram-1-plugin-anatomy")
        diagram_ref = png_file.stem

        # Find corresponding .mmd file
        mmd_file = png_file.with_suffix('.mmd')

        if mmd_file.exists():
            diagrams[diagram_ref] = (str(png_file), str(mmd_file))

    return diagrams


def embed_diagram(doc: Document, diagram_ref: str, diagrams: Dict[str, Tuple[str, str]]) -> None:
    """Embed diagram PNG and Mermaid source into document."""
    if diagram_ref not in diagrams:
        doc.add_paragraph(f"[Diagram not found: {diagram_ref}]")
        return

    png_path, mmd_path = diagrams[diagram_ref]

    # Add diagram title
    p = doc.add_paragraph()
    run = p.add_run(f"Diagram: {diagram_ref}")
    run.font.bold = True
    run.font.size = Pt(12)

    # Add image (max width 6 inches to fit page)
    try:
        doc.add_picture(png_path, width=Inches(6.0))
    except Exception as e:
        doc.add_paragraph(f"[Error loading image: {e}]")
        return

    # Add Mermaid source as code block
    with open(mmd_path, 'r', encoding='utf-8') as f:
        mermaid_source = f.read()

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Mermaid Source:")
    run.font.bold = True

    code_para = doc.add_paragraph(mermaid_source)
    code_para.style = 'Normal'
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()  # Spacing


def process_markdown_line(doc: Document, line: str, diagrams: Dict[str, Tuple[str, str]],
                         in_code_block: bool, code_block_lines: List[str]) -> Tuple[bool, List[str]]:
    """Process a single Markdown line and add to document.

    Returns: (in_code_block, code_block_lines) for state tracking
    """
    # Code block fence
    if line.startswith('```'):
        if in_code_block:
            # End of code block - add it
            if code_block_lines:
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                doc.add_paragraph()  # Spacing
            return False, []
        else:
            # Start of code block
            return True, []

    # Inside code block - accumulate lines
    if in_code_block:
        code_block_lines.append(line)
        return in_code_block, code_block_lines

    # Check for diagram references: "**See diagram**: [6767-f-diagram-1-plugin-anatomy.mmd]"
    # Pattern: **See diagram**: [filename.mmd](filename.mmd) ([PNG version](filename.png))
    diagram_match = re.search(r'\*\*See diagram\*\*:.*?\[(6767-[a-g]-diagram-\d+-[\w-]+)\.mmd\]', line)
    if diagram_match:
        diagram_ref = diagram_match.group(1)
        embed_diagram(doc, diagram_ref, diagrams)
        return in_code_block, code_block_lines

    # Headers
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)

    # Horizontal rule
    elif line.strip() == '---':
        p = doc.add_paragraph()
        p.add_run('_' * 80)

    # Bullet list
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        doc.add_paragraph(text, style='List Bullet')

    # Numbered list (simple detection)
    elif re.match(r'^\d+\.\s+', line):
        text = re.sub(r'^\d+\.\s+', '', line).strip()
        doc.add_paragraph(text, style='List Number')

    # Normal paragraph (skip empty lines)
    elif line.strip():
        # Handle inline formatting
        text = line.strip()

        # Skip metadata lines at top of file
        if text.startswith('**Document Type**:') or text.startswith('**Version**:'):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            doc.add_paragraph(text)

    return in_code_block, code_block_lines


def convert_markdown_to_docx(md_path: Path, output_path: Path, diagrams: Dict[str, Tuple[str, str]]) -> None:
    """Convert a single Markdown file to DOCX with diagram embedding."""
    print(f"Converting {md_path.name} -> {output_path.name}")

    # Create document
    doc = Document()
    setup_document_styles(doc)

    # Read Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Process Markdown line by line
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    for line in lines:
        line = line.rstrip()

        # Table detection (starts with |)
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue
        elif in_table:
            # End of table
            if table_lines:
                add_markdown_table(doc, table_lines)
                doc.add_paragraph()  # Spacing
            in_table = False
            table_lines = []

        # Process line
        in_code_block, code_block_lines = process_markdown_line(
            doc, line, diagrams, in_code_block, code_block_lines
        )

    # Handle final code block if document ended inside one
    if in_code_block and code_block_lines:
        code_text = '\n'.join(code_block_lines)
        p = doc.add_paragraph(code_text)
        p.style = 'Normal'
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # Handle final table if document ended inside one
    if in_table and table_lines:
        add_markdown_table(doc, table_lines)

    # Save document
    doc.save(output_path)
    print(f"  ✓ Created {output_path.name} ({output_path.stat().st_size} bytes)")


def main() -> int:
    """Main entry point."""
    print("=" * 80)
    print("Standard of Truth DOCX Builder")
    print("=" * 80)
    print()

    # Validate project structure
    if not DOCS_DIR.exists():
        print(f"ERROR: Documentation directory not found: {DOCS_DIR}")
        return 1

    success_count = 0

    # Convert each standard
    for standard_name in STANDARDS:
        md_file = DOCS_DIR / f"{standard_name}.md"
        docx_file = DOCS_DIR / f"{standard_name}.docx"

        if not md_file.exists():
            print(f"ERROR: Markdown file not found: {md_file}")
            continue

        # Find diagrams for this standard
        diagrams = find_diagram_files(standard_name)
        if diagrams:
            print(f"Found {len(diagrams)} diagram(s) for {standard_name}")

        try:
            convert_markdown_to_docx(md_file, docx_file, diagrams)
            success_count += 1
        except Exception as e:
            print(f"ERROR converting {standard_name}: {e}")
            import traceback
            traceback.print_exc()

    print()
    print("=" * 80)
    print(f"Conversion complete: {success_count}/{len(STANDARDS)} files created")
    print("=" * 80)

    return 0 if success_count == len(STANDARDS) else 1


if __name__ == '__main__':
    sys.exit(main())
